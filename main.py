import fhir_parser
import tkinter as tk
import tkinter.messagebox
import docx
from docx import Document
import os

fhir = fhir_parser.FHIR('https://localhost:5001/api/', verify_ssl=False)

from tkinter import *
import tkinter.filedialog

observations = fhir.get_patient_observations('b905139e-1601-403c-9d85-f8e3997cdd19')
id=''
window = tk.Tk()
window.title('Form generater')
window.geometry('500x300')
frame_id = Frame(window)
frame_rest = Frame(window)
id_label=Label(frame_id, text="Input Patient ID:")
id_input= Entry(frame_id)
name_text=tk.StringVar()
name_text.set('Name:')
Name_Label=Label(frame_rest, textvariable=name_text)
def form_generation():
    global id
    if id == '':
        messagebox.showwarning("Error", "Please search for a patient first")
        return
    f = tkinter.filedialog.asksaveasfile(filetypes = [('Word document 2007', 'docx')]).name
    os.remove(f)
    if f == None:
        return
    try:
        document = Document()
        document.add_heading("Patient info")
        table = document.add_table(rows=2, cols=4, style='Table Grid')
        patient = fhir.get_patient(id)
        name = patient.name.family+", "+patient.name.given
        address = patient.addresses
        birthDate = patient.birth_date
        maritalStatus = patient.marital_status
        gender = patient.gender
        table.cell(0,0).text="Name:"
        table.cell(0,1).text=name
        table.cell(0,2).text="Gender:"
        table.cell(0,3).text=gender
        table.cell(1,0).text="Birth date:"
        table.cell(1,1).text=str(birthDate)
        table.cell(1,2).text="Marital status:"
        table.cell(1,3).text=str(maritalStatus)
        document.add_heading("Last observation", level = 3)
        observations = fhir.get_patient_observations(id)
        observation = observations.pop()
        observationType = observation.type
        dataTime = observation.issued_datetime
        components = observation.components
        document.add_paragraph("Observation type: "+observation.type)
        table_2 = document.add_table(rows=1, cols=2, style='Table Grid')
        table_2.cell(0,0).text='display'
        table_2.cell(0,1).text='quantity'
        for component in components:
            row=table_2.add_row()
            row.cells[0].text=component.display
            row.cells[1].text=component.quantity()
        if f[-5:].lower()=='.docx':
            f=f[:-5]
        f=f+'.docx'
        document.save(f)
        messagebox.showinfo("Done", "Form generated successfully.")
    except:
        messagebox.showerror("Oops", "Failed generating form.")
confirm_Button=Button(frame_rest, text="Generate Form", command = form_generation)
Name_Label.pack()
confirm_Button.pack(side='bottom')
def id_input_confirm_hit():
    global id
    id = id_input.get()
    if id == '':
        messagebox.showerror(title='Error',message='Please input an ID.')
        return
    try:
        patient = fhir.get_patient(id)
    except:
        id=''
    if id == '':
        messagebox.showerror(title='Error',message='No patient found.')
    else:
        name_text.set('Name: '+patient.full_name())
id_input_confirm = Button(frame_id, text="Confirm", command=id_input_confirm_hit)
id_label.pack(side='left')
id_input.pack(side='left')
id_input_confirm.pack(side='right')
frame_id.pack()
frame_rest.pack(expand=True)
window.mainloop()